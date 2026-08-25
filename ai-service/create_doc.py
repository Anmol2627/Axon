from docx import Document
from docx.shared import Pt
import os

doc = Document()

# Title
title = doc.add_heading('Axon - Project Journey & Documentation', 0)

# Section: The Problem
doc.add_heading('1. The Problem', level=1)
doc.add_paragraph(
    "Assembling the perfect tech team for a project is notoriously slow and error-prone. Project managers often struggle "
    "to accurately define technical requirements, and traditional platforms rely on rigid keyword matching that misses "
    "nuanced expertise. If a project needs 'Next.js', a traditional search might miss an expert who only listed 'React' "
    "and 'Vercel'. This leads to severe skill gaps, delayed projects, and inefficient resource allocation. "
    "Furthermore, reviewing resumes is incredibly time-consuming and often biased."
)

# Section: Why We Chose This
doc.add_heading('2. Why We Chose This Approach', level=1)
doc.add_paragraph(
    "We wanted to build an intelligent intermediary that acts like an expert Technical Recruiter and Software Architect combined. "
    "Instead of traditional keyword searching, we decided to leverage Large Language Models (LLMs) to perform deep semantic "
    "analysis. By extracting structured data from resumes (skills, inferred proficiency, experience) and analyzing project descriptions "
    "(estimating complexity, team size, required roles, and hidden skills), we could mathematically match candidates to projects "
    "using high-dimensional vector search. This ensures candidates are matched on true capability rather than surface-level keywords."
)

# Section: What We Planned vs What Got Implemented
doc.add_heading('3. The Journey: From Plan to Execution', level=1)

doc.add_heading('Phase 1: Foundation & Design (The "Warm Premium" aesthetic)', level=2)
doc.add_paragraph(
    "We started by defining a distinct design philosophy. We explicitly avoided the cold, dark, overly-technical styling "
    "of traditional B2B developer tools (like GitHub or Vercel). Instead, we implemented a 'Warm Premium' design system "
    "characterized by soft off-white backgrounds (Oatmeal), earthy accents (Terracotta), serif typography for headers, "
    "and subtle glassmorphism. We built out the frontend in Next.js (TypeScript) using custom CSS modules to ensure "
    "we had total control over the aesthetics."
)

doc.add_heading('Phase 2: The Core Features Planned', level=2)
doc.add_paragraph("We planned three core pillars for the MVP:")
doc.add_paragraph("1. AI Resume to Profile (Parser)", style='List Bullet')
doc.add_paragraph("2. AI Project Analyzer (Architecture)", style='List Bullet')
doc.add_paragraph("3. Intelligent Candidate Matchmaker (Vector Search)", style='List Bullet')

doc.add_heading('Phase 3: Building the Backend and AI Integration', level=2)
doc.add_paragraph(
    "We realized Next.js API routes would be too restrictive for heavy ML tasks, so we built a separate backend "
    "using FastAPI (Python) and Uvicorn. We integrated Supabase (PostgreSQL) with the pgvector extension to handle "
    "our database and vector embeddings. For the LLM, we initially chose the 'llama3-8b-8192' model via the Groq API "
    "for its speed and JSON formatting capabilities."
)

doc.add_heading('Phase 4: What Broke (The Challenges & Fixes)', level=2)
doc.add_paragraph(
    "During implementation and testing, we ran into several critical issues that threatened the MVP:"
)

p1 = doc.add_paragraph("Model Decommissioning: ", style='List Bullet')
p1.add_run("Suddenly, the Groq API started throwing 404 and 400 errors. We discovered that Groq had deprecated and decommissioned all their Llama 3 models entirely. To fix this, we urgently migrated our entire AI pipeline to the 'qwen/qwen3.6-27b' model.")

p2 = doc.add_paragraph("Thinking Model JSON Parsing: ", style='List Bullet')
p2.add_run("The new Qwen model was a 'reasoning' model that wrapped its output in <think> tags. This broke our strict JSON parser. We had to implement a highly robust Regex parsing system that stripped out Markdown fences, removed <think> blocks, and extracted the raw JSON safely.")

p3 = doc.add_paragraph("Database Schema Mismatches: ", style='List Bullet')
p3.add_run("As the UI evolved to include 'Recommended Workflows', 'Education', and 'Links', the frontend would hang silently because Supabase rejected the payloads (the SQL columns didn't exist). We fixed this by generating SQL migration scripts and adding robust try-catch blocks to the frontend to surface errors immediately.")

p4 = doc.add_paragraph("Python Syntax Errors: ", style='List Bullet')
p4.add_run("While rewriting the LLM prompt to fix the Qwen parsing issue, we accidentally introduced a Python f-string bracket escaping error, which brought the backend down with a 500 Internal Server Error. We debugged the logs and corrected the syntax to restore the parser.")

doc.add_heading('Phase 5: The Final Delivery', level=2)
doc.add_paragraph(
    "After resolving the AI parsing bugs and database migrations, the three core features connected beautifully. "
    "A user could upload a PDF resume, have their skills, experience, and education extracted instantly, create a "
    "project via a simple prompt, and immediately see themselves scored and ranked against the project's requirements. "
    "We successfully pushed the final MVP to GitHub, ready for deployment via Vercel (Frontend) and Render (Backend)."
)

# Save the document
save_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'Axon_Project_Story.docx')
doc.save(save_path)
print(f"Document saved to {save_path}")
